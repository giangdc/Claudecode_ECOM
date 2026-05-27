import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r'E:\AI\Ecom\ecom-pdh\00_input\chucnang_QLdactinh\[Sprint V1.2] URD - Specification Management (1).docx')

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print("="*80)

for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t:
        style = para.style.name
        print(f"P{i:04d} [{style}] {t}")

print("\n" + "="*80)
print("TABLES:")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace('\n',' ') for c in row.cells]
        line = ' | '.join(cells)
        if any(c.strip() for c in cells):
            print(f"  R{ri}: {line[:250]}")
